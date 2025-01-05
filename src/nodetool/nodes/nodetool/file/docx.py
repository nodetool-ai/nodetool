import os
from docx import Document as DocxDocument
from docx.document import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from enum import Enum
from pydantic import Field
from typing import List, Optional
from nodetool.metadata.types import DocumentRef, FilePath, ImageRef
from nodetool.workflows.base_node import BaseNode
from nodetool.workflows.processing_context import ProcessingContext


class ParagraphAlignment(str, Enum):
    LEFT = "LEFT"
    CENTER = "CENTER"
    RIGHT = "RIGHT"
    JUSTIFY = "JUSTIFY"


class CreateDocument(BaseNode):
    """
    Creates a new Word document
    document, docx, file, create
    """

    async def process(self, context: ProcessingContext) -> DocumentRef:
        return DocumentRef(data=DocxDocument())


class LoadWordDocument(BaseNode):
    """
    Loads a Word document from disk
    document, docx, file, load, input
    """

    path: str = Field(default="", description="Path to the document to load")

    async def process(self, context: ProcessingContext) -> DocumentRef:
        if not self.path.strip():
            raise ValueError("path cannot be empty")
        expanded_path = os.path.expanduser(self.path)
        with open(expanded_path, "rb") as f:
            data = DocxDocument(f)
            return DocumentRef(uri="file://" + expanded_path, data=data)


class AddHeading(BaseNode):
    """
    Adds a heading to the document
    document, docx, heading, format
    """

    document: DocumentRef = Field(
        default=DocumentRef(), description="The document to add the heading to"
    )
    text: str = Field(default="", description="The heading text")
    level: int = Field(default=1, ge=1, le=9, description="Heading level (1-9)")

    async def process(self, context: ProcessingContext) -> DocumentRef:
        assert self.document.data is not None, "Document is not connected"
        self.document.data.add_heading(self.text, level=self.level)
        return self.document


class AddParagraph(BaseNode):
    """
    Adds a paragraph of text to the document
    document, docx, text, format
    """

    document: DocumentRef = Field(
        default=DocumentRef(), description="The document to add the paragraph to"
    )
    text: str = Field(default="", description="The paragraph text")
    alignment: ParagraphAlignment = Field(
        default=ParagraphAlignment.LEFT, description="Text alignment"
    )
    bold: bool = Field(default=False, description="Make text bold")
    italic: bool = Field(default=False, description="Make text italic")
    font_size: Optional[int] = Field(default=None, description="Font size in points")

    async def process(self, context: ProcessingContext) -> DocumentRef:
        assert isinstance(self.document.data, Document), "Document is not connected"
        paragraph = self.document.data.add_paragraph()
        run = paragraph.add_run(self.text)

        if self.bold:
            run.bold = True
        if self.italic:
            run.italic = True
        if self.font_size:
            run.font.size = Pt(self.font_size)

        align_map = {
            ParagraphAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            ParagraphAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            ParagraphAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
            ParagraphAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        paragraph.alignment = align_map[self.alignment]

        return self.document


class AddTable(BaseNode):
    """
    Adds a table to the document
    document, docx, table, format
    """

    document: DocumentRef = Field(
        default=DocumentRef(), description="The document to add the table to"
    )
    rows: int = Field(description="Number of rows", ge=1)
    cols: int = Field(description="Number of columns", ge=1)
    data: List[List[str]] = Field(description="Table data as nested lists")

    async def process(self, context: ProcessingContext) -> DocumentRef:
        assert isinstance(self.document.data, Document), "Document is not connected"
        table = self.document.data.add_table(rows=self.rows, cols=self.cols)

        for i, row_data in enumerate(self.data):
            row = table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row[j].text = str(cell_data)

        return self.document


class AddImage(BaseNode):
    """
    Adds an image to the document
    document, docx, image, format
    """

    document: DocumentRef = Field(
        default=DocumentRef(), description="The document to add the image to"
    )
    image: ImageRef = Field(default=ImageRef(), description="The image to add")
    width: Optional[float] = Field(default=None, description="Image width in inches")
    height: Optional[float] = Field(default=None, description="Image height in inches")

    async def process(self, context: ProcessingContext) -> DocumentRef:
        width = Inches(self.width) if self.width else None
        height = Inches(self.height) if self.height else None

        assert isinstance(self.document.data, Document), "Document is not connected"
        image = await context.asset_to_io(self.image)
        self.document.data.add_picture(image, width=width, height=height)
        return self.document


class AddPageBreak(BaseNode):
    """
    Adds a page break to the document
    document, docx, format, layout
    """

    document: DocumentRef = Field(
        default=DocumentRef(), description="The document to add the page break to"
    )

    async def process(self, context: ProcessingContext) -> DocumentRef:
        assert isinstance(self.document.data, Document), "Document is not connected"
        self.document.data.add_page_break()
        return self.document


class SetDocumentProperties(BaseNode):
    """
    Sets document metadata properties
    document, docx, metadata, properties
    """

    document: DocumentRef = Field(
        default=DocumentRef(), description="The document to modify"
    )
    title: str = Field(default="", description="Document title")
    author: str = Field(default="", description="Document author")
    subject: str = Field(default="", description="Document subject")
    keywords: str = Field(default="", description="Document keywords")

    async def process(self, context: ProcessingContext) -> DocumentRef:
        assert isinstance(self.document.data, Document), "Document is not connected"
        core_props = self.document.data.core_properties

        if self.title:
            core_props.title = self.title
        if self.author:
            core_props.author = self.author
        if self.subject:
            core_props.subject = self.subject
        if self.keywords:
            core_props.keywords = self.keywords

        return self.document


class SaveDocument(BaseNode):
    """
    Writes the document to a file
    document, docx, file, save, output
    """

    document: DocumentRef = Field(
        default=DocumentRef(), description="The document to write"
    )
    path: FilePath = Field(
        default=FilePath(), description="The path to write the document to"
    )

    async def process(self, context: ProcessingContext):
        assert isinstance(self.document.data, Document), "Document is not connected"
        assert self.path.path, "Path is not set"
        expanded_path = os.path.expanduser(self.path.path)
        self.document.data.save(expanded_path)
